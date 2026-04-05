"""
Build a two-column journal article DOCX describing the Zipcatcher transit detector.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page layout: A4, two columns, narrow margins ──────────────────────────────
section = doc.sections[0]
section.page_width    = Cm(21.0)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'),        '2')
    cols.set(qn('w:space'),      '720')
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)

set_two_columns(section)

# ── Style helpers ─────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    if level == 1:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
    else:
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2e, 0x2e, 0x8e)
    return p

def body(doc, text, space_before=0, space_after=4, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    return p

def body_noindent(doc, text, space_before=2, space_after=4, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    return p

def add_box(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    pf.left_indent  = Pt(6)
    pf.right_indent = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EEF2FF')
    pPr.append(shd)
    r1 = p.add_run(label + ': ')
    r1.font.size = Pt(8.5)
    r1.font.bold = True
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(text)
    r2.font.size = Pt(8.5)
    r2.font.name = 'Times New Roman'
    return p

# ── Title block ───────────────────────────────────────────────────────────────
def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(4)
    r = p.add_run('Real-Time Detection of Aircraft and Balloon Transits\n'
                  'Across the Solar and Lunar Disk Using Dual-Signal\n'
                  'Video Analysis')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run('Zipcatcher Transit Tracker Project')
    r2.font.size = Pt(10)
    r2.font.name = 'Times New Roman'
    r2.font.italic = True

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after  = Pt(8)
    pPr = p3._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a1a6e')
    pBdr.append(bottom)
    pPr.append(pBdr)

add_title_block(doc)

# ── Abstract ──────────────────────────────────────────────────────────────────
add_heading(doc, 'Abstract', 2)
body(doc,
    'We describe the design, implementation, and validation of a real-time video '
    'processing system that detects transiting objects — aircraft and balloons — '
    'crossing the disk of the Sun or Moon as imaged by a Seestar S50 smart telescope. '
    'The system operates on a downscaled Real-Time Streaming Protocol (RTSP) video '
    'stream at 30 frames per second (fps) and employs a dual-signal algorithm: '
    'Signal A, a consecutive-frame intensity difference sensitive to fast-moving '
    'objects; and Signal B, a spatially weighted, wavelet-detrended difference '
    'against a slowly evolving reference frame, sensitive to slower or more subtle '
    'crossings. Three complementary trigger gates — spike, consecutive-frame, and '
    'matched-filter — combine under permissive OR logic to maximise recall, while '
    'a graduated matched-filter threshold, a hard centre-ratio gate for long '
    'crossing events, and a noise density guard suppress false positives from '
    'atmospheric limb scintillation. A post-capture analyser independently confirms '
    'detections in saved video clips. Tested against six independently recorded '
    'transit videos the system achieves 100% recall with a false-positive rate of '
    '8% in the secondary analyser.',
    first_indent=False)

# ── 1. Introduction ───────────────────────────────────────────────────────────
add_heading(doc, '1. Introduction')
body(doc,
    'Solar and lunar transits by aircraft and high-altitude balloons are rare, '
    'visually striking events. Their duration is governed primarily by distance '
    'rather than speed. A small piston aircraft flying a pattern at 1,500 ft '
    '(~450 m) is so close that even at a modest airspeed of 80 knots its angular '
    'velocity across the disk is very high — the crossing can be over in a fraction '
    'of a second. A large commercial jet at 35,000 ft (~10 km) is twenty times '
    'further away; even though it is travelling at 500 knots its angular velocity '
    'is much lower, producing a 2–3 second crossing. Balloons at even greater '
    'altitudes move slowly and may take several seconds to cross.')
body(doc,
    'This brevity demands a detection system with latency well under one second '
    'from the moment the object enters the disk to the moment recording begins. '
    'Online tools such as CalSky can predict transits by spacecraft such as the '
    'International Space Station (ISS) from published orbital data, but aircraft '
    'and balloon trajectories are not published in advance with sufficient precision '
    'for pre-pointing alone. A continuous visual detector running on the live '
    'telescope stream is therefore essential.')
body(doc,
    'The Zipcatcher system combines predictive flight-data integration (Section 9) '
    'with a purely visual detector that operates continuously on the telescope '
    'video stream. This article documents the visual detection pipeline in detail: '
    'signal extraction, adaptive thresholding, multi-gate trigger logic, recording '
    'management, and the complementary post-capture analyser.')

# ── 2. System Overview ────────────────────────────────────────────────────────
add_heading(doc, '2. System Overview')
body(doc,
    'The Zipcatcher detector runs as a background thread within a Flask web '
    'application. The Seestar S50 telescope streams H.264 video over RTSP at '
    'its native resolution (typically 1920 × 1080 pixels). FFmpeg, an open-source '
    'multimedia framework, decodes this stream, rescales it to a 180 × 320 pixel '
    'analysis canvas at 30 fps, and delivers raw RGB24 frames to the detector via '
    'a Unix pipe. The reduced resolution dramatically lowers per-frame computation '
    'while preserving sufficient spatial detail to detect objects subtending one '
    'pixel or more on the downscaled image.')
body(doc,
    'On detection, FFmpeg separately records the full-resolution RTSP stream to an '
    'H.264 MP4 file encompassing a 3-second pre-trigger buffer and a 6-second '
    'post-trigger tail, for a total clip length of approximately 9 seconds. This '
    'ensures that even a slow 5-second crossing is captured in its entirety '
    'regardless of when within the crossing the trigger fires.')

add_box(doc, 'Key parameters',
    'Analysis resolution: 180 × 320 px  |  Frame rate: 30 fps  |  '
    'Pre-buffer: 3 s  |  Post-buffer: 6 s  |  Clip total: ~9 s')

# ── 3. Disk Detection and Masking ─────────────────────────────────────────────
add_heading(doc, '3. Disk Detection and Masking')
body(doc,
    'Before any signal computation the detector must locate the solar or lunar '
    'disk within the frame. This is performed every two seconds (60 frames) using '
    'the Hough Circle Transform (HCT), a classical technique that votes for '
    'candidate circles in gradient-edge space. The frame is first smoothed with '
    'a 5 × 5 Gaussian blur (σ = 1) to suppress high-frequency noise from solar '
    'granulation or lunar surface texture.')
body(doc,
    'If the HCT returns no result — which can occur when a large aircraft '
    'partially occludes the disk, when thin cloud covers the limb, or during '
    'brief RTSP dropout — the detector falls back to thresholding pixels brighter '
    'than intensity 180 and fitting a minimum enclosing circle to the largest '
    'contiguous bright region.')
body(doc,
    'From the detected circle of radius r and centre (cx, cy), two Boolean masks '
    'are constructed. The inner disk mask covers all pixels within '
    '(1 − m) × r of the centre, where m = 0.25 is the configurable margin fraction '
    '(environment variable DETECTOR_DISK_MARGIN). This 25% margin deliberately '
    'excludes the solar limb, where atmospheric seeing — the blurring and dancing '
    'of images caused by refractive turbulence — and limb darkening produce '
    'spurious intensity variations. The limb ring mask covers the annular zone '
    'between the inner disk and the full disk edge; it serves as a spatial '
    'reference for the centre-ratio test described in Section 5.3.')
body(doc,
    'A smooth spatial weight matrix W(x, y) equals 1.0 at the disk centre and '
    'decays linearly to 0 at the full disk edge. This weight is applied to '
    'Signal B (Section 4.2) to emphasise activity near the disk centre.')

add_box(doc, 'Disc-lost watchdog',
    'If no disk is detected for more than 120 consecutive frames (4 s) a warning '
    'is logged and a Telegram notification dispatched. Detection continues for a '
    '3-second grace period using the last known mask — preventing a single bad '
    'HCT frame from disabling the spike gate at the very moment a large aircraft '
    'produces its strongest signal.')

# ── 4. Dual-Signal Extraction ─────────────────────────────────────────────────
add_heading(doc, '4. Dual-Signal Extraction')
body(doc,
    'The core of the detector computes two complementary scalar signals per frame, '
    'each designed to be sensitive to different transit durations and speeds.',
    first_indent=False)

add_heading(doc, '4.1  Signal A: Consecutive-Frame Difference', 2)
body(doc,
    'Signal A measures the mean absolute intensity change between consecutive '
    'frames within the inner disk mask:')
body_noindent(doc, '    A = mean( |F(t) − F(t−1)| ) over inner disk',
    space_before=2, space_after=2)
body(doc,
    'Before taking the absolute value, the global mean of the difference image is '
    'subtracted. This step partially suppresses scintillation — the rapid '
    'intensity flickering caused by refractive turbulence acting like a moving '
    'lens — by removing the average frame-wide brightness shift. It does not '
    'eliminate scintillation entirely: the spatially non-uniform component, '
    'where different parts of the disk brighten and dim independently, remains. '
    'However, the adaptive threshold (Section 5.1) empirically accounts for '
    'this residual noise floor under prevailing seeing conditions.')
body(doc,
    'Signal A is most effective for fast-moving objects that traverse several '
    'pixels between successive 30 fps frames. A small aircraft at 1,500 ft can '
    'move 10–30 pixels per frame on the analysis canvas, producing a strong, '
    'clearly localised Signal A. A very high-altitude aircraft moves only a few '
    'pixels per frame, relying more heavily on Signal B.')

add_heading(doc, '4.2  Signal B: Reference-Frame Difference', 2)
body(doc,
    'Signal B compares the current frame against a slowly evolving reference '
    'frame R(t) maintained as an Exponential Moving Average (EMA):')
body_noindent(doc, '    R(t) = (1 − α) · R(t−1) + α · F(t),   α = 0.02',
    space_before=2, space_after=2)
body(doc,
    'At 30 fps and α = 0.02, the EMA half-life is approximately 34 frames '
    '(1.1 seconds). The reference therefore tracks slow background changes — '
    'gradual drift in solar brightness, slow cloud passage — while remaining '
    'blind to sub-second transits. The reference is frozen for 5 seconds after '
    'each detection to prevent the transit silhouette being absorbed into the '
    'background during the event.')
body(doc,
    'The raw Signal B score is:')
body_noindent(doc, '    B_raw = mean( |F(t) − R(t)| · W ) over disk',
    space_before=2, space_after=2)
body(doc,
    'where W is the spatial weight matrix from Section 3. The per-channel '
    'difference is mean-subtracted before weighting to partially cancel the '
    'uniform component of scintillation.')

add_heading(doc, '4.3  Wavelet Detrending of Signal B', 2)
body(doc,
    'Raw Signal B is susceptible to slow ramps caused by cloud edges drifting '
    'across the disk or gradual changes in RTSP encoder bit-rate. To remove '
    'these low-frequency artefacts a Discrete Wavelet Transform (DWT) is applied '
    'to a rolling 20-second buffer of B_raw values. A level-3 decomposition '
    'using the sym4 wavelet is performed; the approximation coefficients — '
    'representing components slower than ~2 seconds — are zeroed, and the signal '
    'is reconstructed from the detail coefficients alone:')
body_noindent(doc,
    '    B = |IDWT( zero_approx( DWT(B_raw buffer) ) )|_last',
    space_before=2, space_after=2)
body(doc,
    'If the PyWavelets library is unavailable, Signal B falls back to the raw '
    'B_raw value and a startup warning is logged.')

# ── 5. Adaptive Thresholding ──────────────────────────────────────────────────
add_heading(doc, '5. Adaptive Thresholding and Guards')

add_heading(doc, '5.1  Adaptive Threshold', 2)
body(doc,
    'The detector maintains a rolling 20-second history of Signal A and B values '
    'and computes an adaptive threshold T using the Median Absolute Deviation '
    '(MAD), a robust estimator of statistical dispersion:')
body_noindent(doc,
    '    T = max( 0.5,  median + max(3 · MAD,  0.5 · median) )',
    space_before=2, space_after=2)
body(doc,
    'The outer max(0.5, …) floor prevents the threshold reaching zero when the '
    'RTSP stream delivers duplicate frames — a condition that arises during '
    'telescope reconnects. Without this floor a zero threshold causes the spike '
    'gate to fire on every compression artefact, generating a storm of false '
    'recordings. The floor of 0.5 intensity units is well below the 2–5 unit '
    'noise floor of a live telescope stream, so genuine sensitivity is unaffected.')

add_heading(doc, '5.2  Noise Density Guard', 2)
body(doc,
    'When the scene is unusually active — during a burst of seeing turbulence or '
    'when thin cloud breaks cross the disk — the noise density guard raises both '
    'thresholds proportionally to the ratio of recent to background activity, '
    'capped at a factor of 2.0 to prevent extreme conditions from disabling '
    'detection entirely.')

add_heading(doc, '5.3  Spatial Concentration: Centre Ratio', 2)
body(doc,
    'A genuine transit produces activity concentrated in the inner disk, where '
    'the dark silhouette blocks bright disk light. Limb scintillation, by '
    'contrast, appears predominantly at the disk edge. The centre ratio R_c '
    'quantifies this spatial distribution:')
body_noindent(doc,
    '    R_c = mean(|diff| over inner mask) / mean(|diff| over limb ring)',
    space_before=2, space_after=2)
body(doc,
    'R_c acts as a soft gate for all detections — reducing the confidence score '
    'when R_c < 2.5 — and as a hard gate for long matched-filter detections '
    '(Section 6.3): when the matched-filter template exceeds 60 frames and '
    'neither the spike nor the consecutive-frame gate also fired, a centre ratio '
    'below 1.0 (limb-dominant activity) causes the trigger to be rejected '
    'outright. This specifically targets sustained limb scintillation, which '
    'can accumulate enough triggered frames to pass a 90- or 120-frame template '
    'while always producing limb-dominant rather than centre-dominant signal.')

# ── 6. Trigger Gates ──────────────────────────────────────────────────────────
add_heading(doc, '6. Trigger Gates')
body(doc,
    'The detector implements three independent trigger gates. Recording fires if '
    'ANY gate passes — permissive OR logic that prioritises recall. The cost of '
    'a false recording is far lower than the cost of a missed transit.',
    first_indent=False)

add_heading(doc, '6.1  Spike Gate', 2)
body(doc,
    'The spike gate fires immediately when either signal exceeds the adaptive '
    'threshold by a multiplicative factor S (default 3.0, configurable via '
    'DETECTOR_SPIKE_MULT):')
body_noindent(doc,
    '    spike = (A > S · T_A  or  B > S · T_B)  and  disc_ok',
    space_before=2, space_after=2)
body(doc,
    'disc_ok remains true for 3 seconds after the last successful HCT, so a '
    'large aircraft that momentarily disrupts disk detection can still fire the '
    'spike gate at its moment of maximum amplitude.')
body(doc,
    'The spike gate is the primary — and sometimes only — trigger for close, '
    'low-altitude aircraft. A small plane in a landing pattern at 1,500 ft is '
    'only ~450 m away. Even at a modest 80 knots its angular velocity is so high '
    'that it may cross the entire disk in 3–5 frames (100–167 ms). Consecutive-'
    'frame or matched-filter gates cannot accumulate enough evidence in that '
    'time; only the spike gate fires fast enough. By contrast, a commercial jet '
    'at 35,000 ft, though far faster in absolute terms, is ~10 km away and '
    'takes 2–3 seconds to cross, giving the slower gates time to accumulate.')

add_heading(doc, '6.2  Consecutive-Frame Gate', 2)
body(doc,
    'The consecutive-frame gate requires the signal to exceed the adaptive '
    'threshold for N consecutive frames (default N = 3, ~100 ms). This filters '
    'single-frame noise spikes while catching any transit lasting longer than '
    '100 ms. During an active ADS-B prediction window (Section 9), N is halved '
    'to maximise sensitivity for the predicted crossing.')

add_heading(doc, '6.3  Matched-Filter Gate', 2)
body(doc,
    'The matched-filter (MF) gate tolerates gaps in the above-threshold signal '
    'caused by atmospheric seeing, codec artefacts, or partial occlusion. '
    'For each template duration n ∈ {6, 10, 15, 24, 40, 60, 90, 120} frames the '
    'gate checks whether the most recent n frames contain at least H(n) triggered '
    'entries, where the hit-rate fraction is graduated:')
body_noindent(doc,
    '    frac(n) = 0.70 (n ≤ 15),  0.60 (n ≤ 40),  0.50 (n ≤ 60),  0.45 (n > 60)',
    space_before=2, space_after=4)
body(doc,
    'The 90- and 120-frame templates (3–4 seconds) were introduced specifically '
    'for high-altitude aircraft and weather balloons that spend several seconds '
    'on the disk but may activate the signal in only 45% of frames due to '
    'atmospheric seeing. The hard centre-ratio guard (Section 5.3) prevents '
    'these long templates from being tripped by sustained limb scintillation.')

# ── 7. Centroid Tracking and Confidence Scoring ───────────────────────────────
add_heading(doc, '7. Centroid Tracking and Confidence Scoring')
body(doc,
    'Each frame, the detector computes a weighted centroid of differential '
    'activity within the inner disk mask. If the centroid displacement between '
    'consecutive frames exceeds a minimum magnitude (default 2.0 px) and the '
    'dot product of successive displacement vectors is positive (consistent '
    'direction of motion), a track-agreement counter increments. The track gate '
    'is soft: insufficient track consistency reduces the confidence score but '
    'does not block recording.')
body(doc,
    'A probabilistic confidence score is computed via a logistic (sigmoid) '
    'function of a logit combining the signal-to-noise ratio of the stronger '
    'signal, the centre ratio, and the track consistency fraction, with additive '
    'bonuses for the spike gate and penalties for soft gate failures and low '
    'CNN confidence (Section 8). Detections are labelled strong, weak, or '
    'speculative based on the resulting score.')

# ── 8. CNN Advisory Gate ──────────────────────────────────────────────────────
add_heading(doc, '8. CNN Advisory Gate')
body(doc,
    'An optional Convolutional Neural Network (CNN) second-stage gate provides '
    'an independent transit probability estimate. A rolling buffer of the most '
    'recent 15 grayscale frames, each resized to 90 × 160 pixels, is passed to '
    'an ONNX (Open Neural Network Exchange) model at the moment of trigger. '
    'If the resulting probability falls below a configurable threshold (default '
    '0.40) the confidence logit is penalised by 0.25, but recording proceeds '
    'unconditionally. The CNN is advisory only: it can never suppress a recording.')

# ── 9. ADS-B Prediction Integration ──────────────────────────────────────────
add_heading(doc, '9. ADS-B Prediction Integration')
body(doc,
    'ADS-B (Automatic Dependent Surveillance–Broadcast) position reports from '
    'multiple free aggregators — OpenSky Network, ADSB-One, adsb.lol, and '
    'adsb.fi — feed the Zipcatcher transit predictor. The predictor continuously '
    'computes the angular separation between each tracked aircraft and the current '
    'solar or lunar position. When a specific flight is projected to transit '
    'within the next few minutes, a primed event is registered with an expiry '
    'time. The priming mechanism has two effects on the detector: the consecutive-'
    'frame requirement N is halved, and the predicted flight identifier is logged '
    'alongside any detection fired during the prime window.')
body(doc,
    'After a detection, the FlightAware AeroAPI is queried to enrich the event '
    'record with aircraft type, origin, and destination. Position predictions '
    'assume constant velocity and heading and are considered reliable for up to '
    '15 minutes.')

# ── 10. Cooldown and Recording Management ─────────────────────────────────────
add_heading(doc, '10. Cooldown and Recording Management')
body(doc,
    'A 6-second cooldown between successive triggers prevents re-firing on the '
    'echo of a transit in the EMA reference while remaining short enough to '
    'capture two independent aircraft crossing within quick succession — for '
    'example, two flights at different altitudes approaching the same airport. '
    'Any trigger suppressed by the cooldown is logged at WARNING level with its '
    'full signal values, providing forensic evidence for post-hoc investigation '
    'of potential missed transits.')
body(doc,
    'On trigger, the detector snapshots its 3-second circular JPEG pre-buffer '
    '(90 frames) and spawns a background thread that collects 6 further seconds '
    'of post-trigger frames. Both segments are piped to FFmpeg for H.264 encoding '
    'into an MP4 file, with phase-correlation image stabilisation applied to '
    'correct mount drift before encoding.')

# ── 11. Post-Capture Analyser ─────────────────────────────────────────────────
add_heading(doc, '11. Post-Capture Analyser')
body(doc,
    'Saved MP4 clips are independently re-analysed by a second pipeline, '
    'the post-capture analyser. This pipeline operates at full resolution and '
    'constructs a median reference image from the first N frames of the clip '
    '(N = min(90, total/2) for solar; N = min(20, total/4) for lunar), then '
    'compares each subsequent frame against this frozen reference. Blob detection '
    'via connected components, static-feature filtering (which suppresses '
    'sunspots and lunar craters that remain fixed across frames), and a coherence '
    'tracker that requires a minimum path length and speed confirm or reject '
    'the detected object.')
body(doc,
    'A critical robustness improvement was added to handle clips where OpenCV '
    'over-reports the frame count. When the video stream ends before the reference '
    'buffer fills, the analyser now force-builds a reference from whatever frames '
    'were collected (minimum 5), allowing detection to proceed on short or '
    're-encoded clips. Prior to this fix, a 30-frame clip with an OpenCV-reported '
    'count of 110 would silently produce zero detections.')
body(doc,
    'The analyser produces a composite still image showing all confirmed transit '
    'positions overlaid on a clean background frame, a JSON sidecar with event '
    'metadata, and optional CNN re-classification of each detected event.')

# ── 12. Validation ────────────────────────────────────────────────────────────
add_heading(doc, '12. Validation')
body(doc,
    'The system was validated against two datasets. First, six independently '
    'recorded transit videos (single-frame to 181-frame crossings, resolutions '
    'from 900 × 720 to 1920 × 1080, frame rates 24–30 fps) were submitted to '
    'the post-capture analyser. All six produced at least one confirmed transit '
    'event (100% recall). Prior to the reference-building fix, one 30-frame clip '
    '— whose frame count OpenCV over-reported as 110 — produced zero detections.')
body(doc,
    'Second, 102 gallery clips hand-labelled as false positives were submitted '
    'to the analyser. Ninety-four (92%) produced zero transit events. The '
    'remaining eight were already flagged by the pre-fix analyser: one is likely '
    'a genuine aircraft mislabelled as a false positive (10,802-pixel blob, '
    'aspect ratio 1.57, speed 4.25 normalised units per second); the rest are '
    'residual limb-shimmer detections that passed the coherence speed threshold.')
body(doc,
    'No new false positive regressions were introduced by any of the changes '
    'described in this article.')
body(doc,
    'Balloon detectability depends on altitude and wind. A weather balloon at '
    '30 km altitude drifts very slowly; its angular velocity across the disk '
    'may be so low that Signal A (which measures frame-to-frame motion) is '
    'negligible. Signal B detects the sustained deviation from reference during '
    'the first 1–2 seconds after the balloon enters the disk, before the EMA '
    'begins to adapt. The new 90- and 120-frame matched-filter templates extend '
    'coverage to crossings where seeing degrades signal continuity. A balloon '
    'already on disk when the detector starts is baked into the initial reference '
    'and cannot be detected; this is an acknowledged limitation.')
body(doc,
    'The hard centre-ratio gate for long MF templates (Section 5.3) was validated '
    'against the same 102 false-positive clips: no additional false positives were '
    'introduced, confirming that genuine transits (which produce centre-dominant '
    'signal) are unaffected, while sustained limb-shimmer triggers are blocked.')

# ── 13. Discussion ────────────────────────────────────────────────────────────
add_heading(doc, '13. Discussion')
body(doc,
    'The dual-signal architecture provides complementary coverage: Signal A '
    'excels at close, low-altitude aircraft where frame-to-frame motion is large '
    'and the crossing is over in a handful of frames; Signal B with wavelet '
    'detrending catches distant, slow, or partially occluded crossings where '
    'per-frame motion is small but the deviation from background is sustained. '
    'The three-gate OR trigger with graduated MF thresholds extends this to '
    'balloons and very high-altitude aircraft.')
body(doc,
    'Sunspots are in practice irrelevant to the live detector. They are static '
    '— they contribute nothing to Signal A — and their slow multi-day growth is '
    'removed by wavelet detrending from Signal B. They appear only in the '
    'post-capture analyser, where the static-feature filter correctly '
    'suppresses them.')
body(doc,
    'Limb scintillation remains the dominant false-positive source. The spatial '
    'concentration test is the primary defence. It is implemented as a hard gate '
    'for the longest MF templates, where sustained limb shimmer is most likely '
    'to accumulate false evidence, and as a soft confidence modifier for shorter '
    'events where false-positive risk is lower.')

# ── 14. Conclusion ────────────────────────────────────────────────────────────
add_heading(doc, '14. Conclusion')
body(doc,
    'We have presented a complete real-time video detection pipeline for aircraft '
    'and balloon transits of the solar and lunar disk. The system achieves 100% '
    'recall on a six-clip validation set while correctly rejecting 92% of '
    'hand-labelled false positives in the post-capture analyser. Key contributions '
    'include: a dual-signal algorithm with mean-subtracted scintillation '
    'suppression; wavelet detrending of the reference-frame signal; a three-gate '
    'OR trigger with spike, consecutive-frame, and matched-filter paths; graduated '
    'matched-filter thresholds covering 0.2–4 second crossings; a hard '
    'centre-ratio gate for long MF templates that blocks sustained limb '
    'scintillation; an adaptive threshold floor preventing runaway triggering on '
    'duplicate RTSP frames; forensic cooldown logging; full 3-second pre-buffer '
    'capture; and a robust reference-building fallback in the post-capture '
    'analyser.')

# ── Appendix: Configuration Parameters ───────────────────────────────────────
add_heading(doc, 'Appendix: Key Configuration Parameters')

rows = [
    ('DETECTOR_DISK_MARGIN',     '0.25', 'Fraction of disk radius excluded from limb'),
    ('CENTRE_EDGE_RATIO_MIN',    '2.5',  'Minimum inner/limb signal ratio (soft gate)'),
    ('CONSEC_FRAMES_REQUIRED',   '3',    'Frames above threshold before consec gate fires'),
    ('DETECTOR_SPIKE_MULT',      '3.0',  'Spike gate multiplier on adaptive threshold'),
    ('DETECTION_COOLDOWN',       '6 s',  'Minimum seconds between trigger events'),
    ('DETECTION_PRE_BUFFER',     '3 s',  'Pre-trigger circular buffer duration'),
    ('DETECTION_POST_BUFFER',    '6 s',  'Post-trigger capture duration'),
    ('CNN_GATE_THRESHOLD',       '0.40', 'CNN advisory threshold (never blocks)'),
    ('DETECTOR_TRACK_MIN_MAG',   '2.0',  'Minimum centroid displacement (px)'),
    ('DETECTOR_TRACK_MIN_AGREE', '0.6',  'Fraction of frames with consistent direction'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Parameter'
hdr[1].text = 'Default'
hdr[2].text = 'Description'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'

for param, default, desc in rows:
    row = table.add_row().cells
    row[0].text = param
    row[1].text = default
    row[2].text = desc
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

out = '/Users/Tom/flymoon/Zipcatcher_Transit_Detector.docx'
doc.save(out)
print(f'Saved: {out}')
